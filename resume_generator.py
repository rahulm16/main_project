from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

class ResumeGenerator:
    def __init__(self):
        self.document = Document()
        self._create_styles()
    
    def _create_styles(self):
        # Create custom styles for consistent formatting
        styles = self.document.styles
        
        # Heading 1 style
        h1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.base_style = styles['Heading 1']
        h1_font = h1_style.font
        h1_font.size = Pt(16)
        h1_font.color.rgb = RGBColor(0, 51, 102)
        h1_font.bold = True
        
        # Section style
        section_style = styles.add_style('Section Text', WD_STYLE_TYPE.PARAGRAPH)
        section_font = section_style.font
        section_font.size = Pt(11)
        section_font.name = 'Calibri'
        
        # Header style
        header_style = styles.add_style('Header Name', WD_STYLE_TYPE.PARAGRAPH)
        header_font = header_style.font
        header_font.size = Pt(24)
        header_font.name = 'Calibri'
        header_font.bold = True
        
    def _add_horizontal_line(self):
        paragraph = self.document.add_paragraph()
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '4F81BD')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return paragraph

    def generate_template1(self, user_data):
        # Professional Template
        # Set margins
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Header
        name_paragraph = self.document.add_paragraph(style='Header Name')
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_paragraph.add_run(user_data['name'].upper())
        
        # Contact Info Block
        contact = self.document.add_paragraph(style='Section Text')
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(f"{user_data['email']} \n")
        contact.add_run(user_data.get('github', '\n').replace('https://', ''))
        contact.add_run(f" • {user_data['linkedin'].replace('https://', '')}")
        
        self._add_horizontal_line()
        
        # Professional Summary
        if user_data.get('summary'):
            self.document.add_paragraph('PROFESSIONAL SUMMARY', style='Custom Heading 1')
            summary = self.document.add_paragraph(style='Section Text')
            summary.add_run(user_data['summary'])
            self._add_horizontal_line()
        
        # Education
        self.document.add_paragraph('EDUCATION', style='Custom Heading 1')
        edu = self.document.add_paragraph(style='Section Text')
        edu.add_run(user_data['education']).bold = True
        edu.add_run('\n')
        edu.add_run(f"Specialization: {user_data['specialization']}")
        edu.add_run('\n')
        edu.add_run(f"Course: {user_data['course']}")
        self._add_horizontal_line()
        
        # Skills
        self.document.add_paragraph('TECHNICAL SKILLS', style='Custom Heading 1')
        skills_para = self.document.add_paragraph(style='Section Text')
        
        # Group skills by category if provided
        if isinstance(user_data['skills'], dict):
            for category, skills in user_data['skills'].items():
                skills_para.add_run(f"{category}: ").bold = True
                skills_para.add_run(', '.join(skills))
                skills_para.add_run('\n')
        else:
            skills_para.add_run(', '.join(user_data['skills']))
        
        self._add_horizontal_line()
        
        # Work Experience
        if user_data.get('work_experience'):
            self.document.add_paragraph('PROFESSIONAL EXPERIENCE', style='Custom Heading 1')
            
            if isinstance(user_data['work_experience'], list):
                for experience in user_data['work_experience']:
                    exp_para = self.document.add_paragraph(style='Section Text')
                    exp_para.add_run(experience['title']).bold = True
                    exp_para.add_run(f"\n{experience['company']} • {experience['duration']}")
                    
                    # Add responsibilities as bullet points
                    for responsibility in experience['responsibilities']:
                        bullet_para = self.document.add_paragraph(style='Section Text')
                        bullet_para.style = 'List Bullet'
                        bullet_para.add_run(responsibility)
            else:
                exp_para = self.document.add_paragraph(style='Section Text')
                exp_para.add_run(user_data['work_experience'])
            
            self._add_horizontal_line()
        
        # Projects (if available)
        if user_data.get('projects'):
            self.document.add_paragraph('PROJECTS', style='Custom Heading 1')
            for project in user_data['projects']:
                proj_para = self.document.add_paragraph(style='Section Text')
                proj_para.add_run(project['name']).bold = True
                proj_para.add_run(f"\n{project['description']}")
                if project.get('technologies'):
                    proj_para.add_run(f"\nTechnologies used: {', '.join(project['technologies'])}")
                proj_para.add_run('\n')
            
            self._add_horizontal_line()
        
        # Certifications (if available)
        if user_data.get('certifications'):
            self.document.add_paragraph('CERTIFICATIONS', style='Custom Heading 1')
            cert_para = self.document.add_paragraph(style='Section Text')
            for cert in user_data['certifications']:
                cert_para.add_run(f"• {cert}\n")
            
            self._add_horizontal_line()
        
        # Languages (if available)
        if user_data.get('languages'):
            self.document.add_paragraph('LANGUAGES', style='Custom Heading 1')
            lang_para = self.document.add_paragraph(style='Section Text')
            lang_para.add_run(', '.join(user_data['languages']))
        
        return self.document
    
    def generate_template2(self, user_data):
        """
        Modern template with enhanced styling, better spacing, and professional formatting
        """
        self.document = Document()
        sections = self.document.sections
        
        # Set custom page margins
        for section in sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Define custom styles
        styles = self.document.styles
        
        # Name style
        name_style = styles.add_style('Name Style', WD_STYLE_TYPE.PARAGRAPH)
        name_style.font.name = 'Calibri'
        name_style.font.size = Pt(28)
        name_style.font.bold = True
        name_style.font.color.rgb = RGBColor(0, 0, 0)
        name_style.paragraph_format.space_after = Pt(0)
        
        # Title style
        title_style = styles.add_style('Title Style', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(14)
        title_style.font.color.rgb = RGBColor(89, 89, 89)
        title_style.paragraph_format.space_before = Pt(0)
        title_style.paragraph_format.space_after = Pt(12)
        
        # Section heading style
        heading_style = styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(16)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(44, 116, 181)  # Professional blue
        heading_style.paragraph_format.space_before = Pt(16)
        heading_style.paragraph_format.space_after = Pt(8)
        
        # Body text style
        body_style = styles.add_style('Body Style', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Calibri'
        body_style.font.size = Pt(11)
        body_style.font.color.rgb = RGBColor(0, 0, 0)
        body_style.paragraph_format.space_after = Pt(6)
        
        # Header Section
        name_para = self.document.add_paragraph(style='Name Style')
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Add this line to center-align the name
        name_para.add_run(user_data['name'].upper())
        
        if user_data.get('title'):
            title_para = self.document.add_paragraph(style='Title Style')
            title_para.add_run(user_data['title'])
        
        # Contact Information with horizontal layout
        contact_table = self.document.add_table(rows=1, cols=3)  # Changed to 3 columns
        contact_table.autofit = False
        contact_table.allow_autofit = False
        
        # Set equal column widths
        for cell in contact_table.columns:
            cell.width = Inches(2.5)  # Adjusted width for 3 columns
        
        contact_cells = contact_table.rows[0].cells
        
        # Add contact information with icons (using Unicode)
        contact_info = [
            ("✉", user_data['email']),
            ("🔗", user_data.get('github', '').replace('https://', '')), 
            ("💼", user_data['linkedin'].replace('https://', ''))
        ]
        
        for i, (icon, info) in enumerate(contact_info):
            if info:
                p = contact_cells[i].paragraphs[0]
                p.style = body_style
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(f"{icon} {info}")
        
        self.document.add_paragraph()  # Add spacing
        
        # Create main content table
        main_table = self.document.add_table(rows=1, cols=2)
        main_table.allow_autofit = False
        
        # Set column widths (35% and 65%)
        main_table.columns[0].width = Inches(2.8)
        main_table.columns[1].width = Inches(5.2)
        
        left_cell, right_cell = main_table.rows[0].cells
        
        # Add subtle borders to main table
        self._set_cell_borders(left_cell, right=True)
        self._set_cell_borders(right_cell, left=True)
        
        # Left Column Content
        self._add_section(left_cell, "EXPERTISE", user_data.get('skills', []))
        self._add_section(left_cell, "EDUCATION", [
            f"• {user_data['education']}",
            f"• {user_data['specialization']}",
            f"• {user_data['course']}"
        ])
        
        if user_data.get('languages'):
            self._add_section(left_cell, "LANGUAGES", [f"• {lang}" for lang in user_data['languages']])
        
        if user_data.get('certifications'):
            self._add_section(left_cell, "CERTIFICATIONS", [f"• {cert}" for cert in user_data['certifications']])
        
        # Right Column Content
        if user_data.get('summary'):
            self._add_section(right_cell, "PROFESSIONAL SUMMARY", [user_data['summary']], add_bullets=False)
        
        # Work Experience Section - Updated Implementation
        if user_data.get('work_experience'):
            work_heading = right_cell.add_paragraph("PROFESSIONAL EXPERIENCE", style='Section Heading')
            
            if isinstance(user_data['work_experience'], list):
                for exp in user_data['work_experience']:
                    # Company and Title
                    exp_para = right_cell.add_paragraph(style='Body Style')
                    exp_para.add_run(exp['title']).bold = True
                    exp_para.add_run(f"\n{exp['company']}")
                    
                    # Duration
                    dur_para = right_cell.add_paragraph(style='Body Style')
                    dur_para.add_run(exp['duration']).italic = True
                    
                    # Responsibilities
                    if 'responsibilities' in exp:
                        for resp in exp['responsibilities']:
                            resp_para = right_cell.add_paragraph(style='Body Style')
                            resp_para.style = 'List Bullet'
                            resp_para.add_run(resp)
                    
                    right_cell.add_paragraph()  # Spacing between experiences
            else:
                # Handle single string work experience
                exp_para = right_cell.add_paragraph(style='Body Style')
                exp_para.add_run(user_data['work_experience'])
        
        # Projects Section
        if user_data.get('projects'):
            self._add_section(right_cell, "KEY PROJECTS", [], add_bullets=False)
            
            for project in user_data['projects']:
                # Project name and description
                p = right_cell.add_paragraph(style='Body Style')
                p.add_run(project['name']).bold = True
                p.add_run(f"\n{project['description']}")
                
                # Technologies used
                if project.get('technologies'):
                    tech_para = right_cell.add_paragraph(style='Body Style')
                    tech_para.add_run("Technologies: ").italic = True
                    tech_para.add_run(', '.join(project['technologies']))
                
                right_cell.add_paragraph()  # Add spacing between projects
        
        return self.document

    def _add_section(self, cell, heading, items, add_bullets=True):
        """Enhanced helper method to add sections with better formatting"""
        heading_para = cell.add_paragraph(style='Section Heading')
        heading_para.add_run(heading)
        
        if isinstance(items, dict):
            for category, skills in items.items():
                p = cell.add_paragraph(style='Body Style')
                p.add_run(f"{category}\n").bold = True
                skills_list = [f"• {skill}" for skill in skills] if add_bullets else skills
                p.add_run('\n'.join(skills_list))
        elif isinstance(items, list):
            for item in items:
                p = cell.add_paragraph(style='Body Style')
                p.add_run(item)
        else:
            p = cell.add_paragraph(style='Body Style')
            p.add_run(items)

    def _set_cell_borders(self, cell, **kwargs):
        """Helper method to set cell borders"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        for edge, val in kwargs.items():
            if val:
                border = parse_xml(f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                f'<w:{edge} w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                                f'</w:tcBorders>')
                tcPr.append(border)
    
    def generate_template3(self, user_data):
        """
        Creative template with a modern, distinctive design featuring accent colors
        and creative section layouts
        """
        self.document = Document()
        sections = self.document.sections
        
        # Set document margins
        for section in sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # Create custom styles for creative template
        styles = self.document.styles
        
        # Creative heading style
        creative_heading = styles.add_style('Creative Heading', WD_STYLE_TYPE.PARAGRAPH)
        creative_heading.font.size = Pt(16)
        creative_heading.font.color.rgb = RGBColor(51, 122, 183)  # Blue accent color
        creative_heading.font.bold = True
        creative_heading.font.name = 'Helvetica'
        
        # Creative subheading style
        creative_subheading = styles.add_style('Creative Subheading', WD_STYLE_TYPE.PARAGRAPH)
        creative_subheading.font.size = Pt(12)
        creative_subheading.font.color.rgb = RGBColor(68, 68, 68)
        creative_subheading.font.bold = True
        creative_subheading.font.name = 'Helvetica'
        
        # Creative text style
        creative_text = styles.add_style('Creative Text', WD_STYLE_TYPE.PARAGRAPH)
        creative_text.font.size = Pt(11)
        creative_text.font.name = 'Helvetica'
        creative_text.paragraph_format.space_after = Pt(6)
        
        # Header section with distinctive design
        header_table = self.document.add_table(rows=2, cols=1)
        header_table.style = 'Table Grid'
        header_table.autofit = False
        header_cells = header_table.rows[0].cells
        
        # Name and title cell with background
        name_cell = header_cells[0]
        name_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E6E6FA"/>'.format(nsdecls('w'))))  # Light purple background
        
        name_paragraph = name_cell.paragraphs[0]
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_paragraph.add_run(user_data['name'].upper())
        name_run.font.size = Pt(28)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(51, 51, 51)
        
        if user_data.get('title'):
            title_run = name_paragraph.add_run(f"\n{user_data['title']}")
            title_run.font.size = Pt(16)
            title_run.font.color.rgb = RGBColor(102, 102, 102)
        
        # Contact information with icons (using unicode characters)
        contact_cell = header_table.rows[1].cells[0]
        contact_paragraph = contact_cell.paragraphs[0]
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.style = creative_text
        
        contact_info = [
            ("✉", user_data['email']),
            ("🔗", user_data.get('github', '').replace('https://', '')), 
            ("💼", user_data['linkedin'].replace('https://', ''))
        ]
        
        for icon, info in contact_info:
            if info:
                contact_run = contact_paragraph.add_run(f"{icon} {info}   ")
                contact_run.font.size = Pt(10)
        
        self.document.add_paragraph()  # Spacing
        
        # Professional Summary with accent border
        if user_data.get('summary'):
            summary_table = self.document.add_table(rows=1, cols=1)
            summary_cell = summary_table.rows[0].cells[0]
            
            # Fixed XML with proper namespace declaration
            tcBorders = parse_xml(f'''<w:tcBorders {nsdecls('w')}>
                <w:left w:val="single" w:sz="24" w:space="0" w:color="517AB7"/>
            </w:tcBorders>''')
            
            summary_cell._tc.get_or_add_tcPr().append(tcBorders)
            
            summary_para = summary_cell.paragraphs[0]
            summary_para.add_run("PROFESSIONAL SUMMARY").bold = True
            summary_para.add_run(f"\n{user_data['summary']}")
            
            self.document.add_paragraph()
        
        # Skills section with creative layout
        if user_data.get('skills'):
            skills_heading = self.document.add_paragraph("TECHNICAL EXPERTISE", style='Creative Heading')
            skills_table = self.document.add_table(rows=1, cols=3)
            
            if isinstance(user_data['skills'], dict):
                cell_idx = 0
                for category, skills in user_data['skills'].items():
                    cell = skills_table.rows[0].cells[cell_idx % 3]
                    p = cell.add_paragraph(style='Creative Text')
                    p.add_run(f"{category}\n").bold = True
                    p.add_run("• " + "\n• ".join(skills))
                    cell_idx += 1
            else:
                cell = skills_table.rows[0].cells[0]
                p = cell.add_paragraph(style='Creative Text')
                p.add_run("• " + "\n• ".join(user_data['skills']))
            
            self.document.add_paragraph()
        
        # Experience section with timeline-style layout
        if user_data.get('work_experience'):
            self.document.add_paragraph("PROFESSIONAL JOURNEY", style='Creative Heading')
            
            if isinstance(user_data['work_experience'], list):
                for exp in user_data['work_experience']:
                    exp_table = self.document.add_table(rows=1, cols=2)
                    exp_table.allow_autofit = False
                    
                    # Set column widths (30% - 70%)
                    exp_table.columns[0].width = Inches(2)
                    exp_table.columns[1].width = Inches(5)
                    
                    # Timeline column
                    time_cell = exp_table.rows[0].cells[0]
                    time_para = time_cell.add_paragraph(style='Creative Text')
                    time_para.add_run(exp['duration']).bold = True
                    
                    # Details column
                    details_cell = exp_table.rows[0].cells[1]
                    title_para = details_cell.add_paragraph(style='Creative Subheading')
                    title_para.add_run(exp['title'])
                    company_para = details_cell.add_paragraph(style='Creative Text')
                    company_para.add_run(exp['company'])
                    
                    # Responsibilities
                    for resp in exp['responsibilities']:
                        resp_para = details_cell.add_paragraph(style='Creative Text')
                        resp_para.style = 'List Bullet'
                        resp_para.add_run(resp)
                    
                    self.document.add_paragraph()
            else:
                self.document.add_paragraph(user_data['work_experience'], style='Creative Text')
        
        # Projects section with card-style layout
        if user_data.get('projects'):
            self.document.add_paragraph("FEATURED PROJECTS", style='Creative Heading')
            
            for project in user_data['projects']:
                project_table = self.document.add_table(rows=1, cols=1)
                project_cell = project_table.rows[0].cells[0]
                
                # Add subtle background to project cards
                project_cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w'))))
                
                name_para = project_cell.add_paragraph(style='Creative Subheading')
                name_para.add_run(project['name'])
                
                desc_para = project_cell.add_paragraph(style='Creative Text')
                desc_para.add_run(project['description'])
                
                if project.get('technologies'):
                    tech_para = project_cell.add_paragraph(style='Creative Text')
                    tech_para.add_run("Technologies: ").bold = True
                    tech_para.add_run(', '.join(project['technologies']))
                
                self.document.add_paragraph()
        
        # Education and Certifications in a side-by-side layout
        edu_cert_table = self.document.add_table(rows=1, cols=2)
        
        # Education column
        edu_cell = edu_cert_table.rows[0].cells[0]
        edu_cell.add_paragraph("EDUCATION", style='Creative Heading')
        edu_para = edu_cell.add_paragraph(style='Creative Text')
        edu_para.add_run(user_data['education']).bold = True
        edu_para.add_run(f"\nSpecialization: {user_data['specialization']}")
        edu_para.add_run(f"\nCourse: {user_data['course']}")
        
        # Certifications column
        if user_data.get('certifications'):
            cert_cell = edu_cert_table.rows[0].cells[1]
            cert_cell.add_paragraph("CERTIFICATIONS", style='Creative Heading')
            for cert in user_data['certifications']:
                cert_para = cert_cell.add_paragraph(style='Creative Text')
                cert_para.add_run(f"• {cert}")
        
        return self.document
    def save_document(self, filename):
        self.document.save(filename)
